"""
Генерация документов DOCX из шаблонов по заполненным полям.
Возвращает готовый файл в base64 для скачивания.
"""
import json
import os
import base64
import io
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_OK = True
except ImportError:
    DOCX_OK = False


def handler(event: dict, context) -> dict:
    headers = {
        'Access-Control-Allow-Origin': '*',
        'Access-Control-Allow-Methods': 'POST, OPTIONS',
        'Access-Control-Allow-Headers': 'Content-Type, X-Auth-Token',
        'Content-Type': 'application/json'
    }

    if event.get('httpMethod') == 'OPTIONS':
        return {'statusCode': 200, 'headers': headers, 'body': ''}

    body = {}
    if event.get('body'):
        body = json.loads(event['body'])

    template_id = body.get('template_id')
    fields = body.get('fields', {})

    if not template_id:
        return {'statusCode': 400, 'headers': headers, 'body': json.dumps({'error': 'template_id required'})}

    if not DOCX_OK:
        return {'statusCode': 500, 'headers': headers, 'body': json.dumps({'error': 'python-docx not installed'})}

    today = date.today().strftime('%d.%m.%Y')
    generators = {
        1: gen_supply_contract,
        2: gen_rent_contract,
        3: gen_employment_contract,
        4: gen_act,
        5: gen_invoice,
        6: gen_commercial_offer,
        7: gen_power_of_attorney,
        8: gen_nda,
        9: gen_order,
    }

    gen_fn = generators.get(int(template_id))
    if not gen_fn:
        return {'statusCode': 404, 'headers': headers, 'body': json.dumps({'error': 'Template not found'})}

    doc, filename = gen_fn(fields, today)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    file_b64 = base64.b64encode(buf.read()).decode()

    return {
        'statusCode': 200,
        'headers': headers,
        'body': json.dumps({'file_data': file_b64, 'file_name': filename, 'mime': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'})
    }


def base_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)
    return doc


def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    return p


def para(doc, text='', bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
    p.add_run(text)
    return p


def field(fields, key, default='________________'):
    v = fields.get(key, '').strip()
    return v if v else default


# ─── 1. Договор поставки ─────────────────────────────────────────
def gen_supply_contract(fields, today):
    doc = base_doc()
    city = field(fields, 'city', 'г. Москва')
    num = field(fields, 'contract_num', '___')
    d = field(fields, 'contract_date', today)
    supplier = field(fields, 'supplier', '____________________')
    buyer = field(fields, 'buyer', '____________________')
    goods = field(fields, 'goods', '____________________')
    amount = field(fields, 'amount', '____________________')
    delivery_days = field(fields, 'delivery_days', '30')
    supplier_inn = field(fields, 'supplier_inn', '________________')
    buyer_inn = field(fields, 'buyer_inn', '________________')

    heading(doc, 'ДОГОВОР ПОСТАВКИ')
    heading(doc, f'№ {num}', 2)
    p = doc.add_paragraph()
    p.add_run(f'{city}')
    p.add_run(f'\t\t\t\t\t\t\t\t{d}')

    doc.add_paragraph()
    para(doc, f'{supplier}, именуемое в дальнейшем «Поставщик», и {buyer}, именуемое в дальнейшем «Покупатель», заключили настоящий договор о нижеследующем:')

    heading(doc, '1. ПРЕДМЕТ ДОГОВОРА', 2)
    para(doc, f'1.1. Поставщик обязуется поставить, а Покупатель принять и оплатить товар: {goods}.')
    para(doc, f'1.2. Срок поставки: {delivery_days} календарных дней с момента подписания договора.')

    heading(doc, '2. ЦЕНА И ПОРЯДОК РАСЧЁТОВ', 2)
    para(doc, f'2.1. Общая стоимость товара составляет: {amount} рублей, в том числе НДС 20%.')
    para(doc, '2.2. Оплата производится в течение 5 банковских дней с момента получения счёта.')

    heading(doc, '3. КАЧЕСТВО ТОВАРА', 2)
    para(doc, '3.1. Качество поставляемого товара должно соответствовать требованиям действующих ГОСТ, ТУ и иных нормативных документов.')
    para(doc, '3.2. Поставщик предоставляет сертификаты качества и иные документы, подтверждающие качество товара.')

    heading(doc, '4. ОТВЕТСТВЕННОСТЬ СТОРОН', 2)
    para(doc, '4.1. За нарушение сроков поставки Поставщик уплачивает пеню в размере 0,1% от стоимости непоставленного товара за каждый день просрочки.')
    para(doc, '4.2. За нарушение сроков оплаты Покупатель уплачивает пеню в размере 0,1% от неоплаченной суммы за каждый день просрочки.')

    heading(doc, '5. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН', 2)
    p = doc.add_paragraph()
    p.add_run('ПОСТАВЩИК:\n').bold = True
    p.add_run(f'{supplier}\nИНН: {supplier_inn}\n\nПодпись: _______________  /_______________/')

    p = doc.add_paragraph()
    p.add_run('ПОКУПАТЕЛЬ:\n').bold = True
    p.add_run(f'{buyer}\nИНН: {buyer_inn}\n\nПодпись: _______________  /_______________/')

    return doc, f'Договор_поставки_{num}.docx'


# ─── 2. Договор аренды ───────────────────────────────────────────
def gen_rent_contract(fields, today):
    doc = base_doc()
    city = field(fields, 'city', 'г. Москва')
    num = field(fields, 'contract_num', '___')
    d = field(fields, 'contract_date', today)
    landlord = field(fields, 'landlord', '____________________')
    tenant = field(fields, 'tenant', '____________________')
    address = field(fields, 'address', '____________________')
    area = field(fields, 'area', '___')
    rent = field(fields, 'rent', '____________________')
    period_from = field(fields, 'period_from', '01.01.2025')
    period_to = field(fields, 'period_to', '31.12.2025')

    heading(doc, 'ДОГОВОР АРЕНДЫ ПОМЕЩЕНИЯ')
    heading(doc, f'№ {num}', 2)
    p = doc.add_paragraph(); p.add_run(f'{city}\t\t\t\t\t\t\t\t{d}')
    doc.add_paragraph()
    para(doc, f'{landlord}, именуемый в дальнейшем «Арендодатель», и {tenant}, именуемый «Арендатор», заключили настоящий договор:')

    heading(doc, '1. ПРЕДМЕТ ДОГОВОРА', 2)
    para(doc, f'1.1. Арендодатель передаёт во временное пользование нежилое помещение по адресу: {address}, общей площадью {area} кв.м.')
    para(doc, f'1.2. Срок аренды: с {period_from} по {period_to}.')

    heading(doc, '2. АРЕНДНАЯ ПЛАТА', 2)
    para(doc, f'2.1. Ежемесячная арендная плата составляет {rent} рублей, в том числе НДС 20%.')
    para(doc, '2.2. Оплата производится не позднее 5-го числа каждого месяца.')
    para(doc, '2.3. Коммунальные услуги оплачиваются Арендатором самостоятельно.')

    heading(doc, '3. ПРАВА И ОБЯЗАННОСТИ СТОРОН', 2)
    para(doc, '3.1. Арендодатель обязан передать помещение в состоянии, пригодном для использования.')
    para(doc, '3.2. Арендатор обязан использовать помещение по назначению, поддерживать его в надлежащем состоянии.')
    para(doc, '3.3. Арендатор не вправе сдавать помещение в субаренду без письменного согласия Арендодателя.')

    heading(doc, '4. РАСТОРЖЕНИЕ ДОГОВОРА', 2)
    para(doc, '4.1. Договор может быть расторгнут по соглашению сторон.')
    para(doc, '4.2. Каждая из сторон вправе расторгнуть договор, уведомив другую сторону за 30 дней.')

    heading(doc, '5. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН', 2)
    p = doc.add_paragraph(); p.add_run('АРЕНДОДАТЕЛЬ:\n').bold = True; p.add_run(f'{landlord}\n\nПодпись: _______________  /_______________/')
    p = doc.add_paragraph(); p.add_run('АРЕНДАТОР:\n').bold = True; p.add_run(f'{tenant}\n\nПодпись: _______________  /_______________/')

    return doc, f'Договор_аренды_{num}.docx'


# ─── 3. Трудовой договор ─────────────────────────────────────────
def gen_employment_contract(fields, today):
    doc = base_doc()
    city = field(fields, 'city', 'г. Москва')
    num = field(fields, 'contract_num', '___')
    d = field(fields, 'contract_date', today)
    employer = field(fields, 'employer', '____________________')
    employee = field(fields, 'employee', '____________________')
    position = field(fields, 'position', '____________________')
    salary = field(fields, 'salary', '____________________')
    start_date = field(fields, 'start_date', today)
    schedule = field(fields, 'schedule', '5/2, 09:00–18:00')

    heading(doc, 'ТРУДОВОЙ ДОГОВОР')
    heading(doc, f'№ {num}', 2)
    p = doc.add_paragraph(); p.add_run(f'{city}\t\t\t\t\t\t\t\t{d}')
    doc.add_paragraph()
    para(doc, f'{employer}, именуемый в дальнейшем «Работодатель», и {employee}, именуемый «Работник», заключили настоящий трудовой договор:')

    heading(doc, '1. ПРЕДМЕТ ДОГОВОРА', 2)
    para(doc, f'1.1. Работодатель принимает Работника на должность: {position}.')
    para(doc, f'1.2. Дата начала работы: {start_date}.')
    para(doc, '1.3. Договор является бессрочным.')

    heading(doc, '2. ПРАВА И ОБЯЗАННОСТИ РАБОТНИКА', 2)
    para(doc, '2.1. Работник обязан добросовестно исполнять трудовые обязанности, соблюдать правила внутреннего трудового распорядка.')
    para(doc, '2.2. Работник имеет право на предоставление ежегодного оплачиваемого отпуска продолжительностью 28 календарных дней.')

    heading(doc, '3. РАБОЧЕЕ ВРЕМЯ И ОПЛАТА ТРУДА', 2)
    para(doc, f'3.1. Режим работы: {schedule}.')
    para(doc, f'3.2. Должностной оклад: {salary} рублей в месяц до вычета НДФЛ.')
    para(doc, '3.3. Выплата заработной платы производится 10-го и 25-го числа каждого месяца.')

    heading(doc, '4. ГАРАНТИИ И КОМПЕНСАЦИИ', 2)
    para(doc, '4.1. На Работника распространяются все гарантии, предусмотренные Трудовым кодексом РФ.')
    para(doc, '4.2. Работодатель осуществляет обязательное социальное страхование Работника.')

    heading(doc, '5. РЕКВИЗИТЫ И ПОДПИСИ СТОРОН', 2)
    p = doc.add_paragraph(); p.add_run('РАБОТОДАТЕЛЬ:\n').bold = True; p.add_run(f'{employer}\n\nПодпись: _______________  /_______________/')
    p = doc.add_paragraph(); p.add_run('РАБОТНИК:\n').bold = True; p.add_run(f'{employee}\n\nПодпись: _______________  /_______________/')

    return doc, f'Трудовой_договор_{num}.docx'


# ─── 4. Акт выполненных работ ────────────────────────────────────
def gen_act(fields, today):
    doc = base_doc()
    city = field(fields, 'city', 'г. Москва')
    num = field(fields, 'act_num', '___')
    d = field(fields, 'act_date', today)
    contractor = field(fields, 'contractor', '____________________')
    customer = field(fields, 'customer', '____________________')
    work_desc = field(fields, 'work_desc', '____________________')
    amount = field(fields, 'amount', '____________________')
    contract_num = field(fields, 'contract_num', '___')
    contract_date = field(fields, 'contract_date', '___')

    heading(doc, 'АКТ ВЫПОЛНЕННЫХ РАБОТ (ОКАЗАННЫХ УСЛУГ)')
    heading(doc, f'№ {num} от {d}', 2)
    p = doc.add_paragraph(); p.add_run(f'{city}\t\t\t\t\t\t\t\t{d}')
    doc.add_paragraph()
    para(doc, f'{contractor}, именуемый «Исполнитель», и {customer}, именуемый «Заказчик», составили настоящий акт о нижеследующем:')
    doc.add_paragraph()
    para(doc, f'В соответствии с договором № {contract_num} от {contract_date} Исполнитель выполнил, а Заказчик принял следующие работы (услуги):')
    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=4)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '№'; hdr[1].text = 'Наименование работ/услуг'; hdr[2].text = 'Ед. изм.'; hdr[3].text = 'Сумма, руб.'
    for cell in hdr: cell.paragraphs[0].runs[0].bold = True
    r1 = table.rows[1].cells
    r1[0].text = '1'; r1[1].text = work_desc; r1[2].text = 'шт.'; r1[3].text = amount
    r2 = table.rows[2].cells
    r2[1].text = 'ИТОГО:'; r2[3].text = f'{amount} руб.'
    r2[1].paragraphs[0].runs[0].bold = True; r2[3].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    para(doc, f'Общая стоимость выполненных работ составляет: {amount} рублей, в том числе НДС 20%.')
    para(doc, 'Работы выполнены в полном объёме и в установленные сроки. Заказчик претензий не имеет.')
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run('ИСПОЛНИТЕЛЬ:\n').bold = True; p.add_run(f'{contractor}\n\nПодпись: _______________  /_______________/')
    p = doc.add_paragraph(); p.add_run('ЗАКАЗЧИК:\n').bold = True; p.add_run(f'{customer}\n\nПодпись: _______________  /_______________/')

    return doc, f'Акт_выполненных_работ_{num}.docx'


# ─── 5. Счёт на оплату ───────────────────────────────────────────
def gen_invoice(fields, today):
    doc = base_doc()
    num = field(fields, 'invoice_num', '___')
    d = field(fields, 'invoice_date', today)
    seller = field(fields, 'seller', '____________________')
    seller_inn = field(fields, 'seller_inn', '________________')
    seller_bank = field(fields, 'seller_bank', '____________________')
    seller_account = field(fields, 'seller_account', '____________________')
    seller_bik = field(fields, 'seller_bik', '________________')
    buyer = field(fields, 'buyer', '____________________')
    item_name = field(fields, 'item_name', '____________________')
    qty = field(fields, 'qty', '1')
    price = field(fields, 'price', '____________________')
    total = field(fields, 'total', '____________________')

    heading(doc, f'СЧЁТ НА ОПЛАТУ № {num} от {d}')
    doc.add_paragraph()

    p = doc.add_paragraph(); p.add_run('Поставщик (Исполнитель): ').bold = True; p.add_run(seller)
    p = doc.add_paragraph(); p.add_run('ИНН: ').bold = True; p.add_run(seller_inn)
    p = doc.add_paragraph(); p.add_run('Банк: ').bold = True; p.add_run(seller_bank)
    p = doc.add_paragraph(); p.add_run('Р/с: ').bold = True; p.add_run(seller_account)
    p = doc.add_paragraph(); p.add_run('БИК: ').bold = True; p.add_run(seller_bik)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('Покупатель (Плательщик): ').bold = True; p.add_run(buyer)
    doc.add_paragraph()

    table = doc.add_table(rows=3, cols=5)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = '№'; hdr[1].text = 'Наименование'; hdr[2].text = 'Кол-во'; hdr[3].text = 'Цена, руб.'; hdr[4].text = 'Сумма, руб.'
    for cell in hdr: cell.paragraphs[0].runs[0].bold = True
    r1 = table.rows[1].cells
    r1[0].text = '1'; r1[1].text = item_name; r1[2].text = qty; r1[3].text = price; r1[4].text = total
    r2 = table.rows[2].cells
    r2[3].text = 'ИТОГО:'; r2[4].text = f'{total} руб.'
    r2[3].paragraphs[0].runs[0].bold = True; r2[4].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    para(doc, f'Итого к оплате: {total} рублей, в том числе НДС 20%.')
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('Руководитель: _______________  /_______________/')
    p = doc.add_paragraph(); p.add_run('Главный бухгалтер: _______________  /_______________/')

    return doc, f'Счёт_{num}.docx'


# ─── 6. Коммерческое предложение ─────────────────────────────────
def gen_commercial_offer(fields, today):
    doc = base_doc()
    d = field(fields, 'offer_date', today)
    sender = field(fields, 'sender', '____________________')
    recipient = field(fields, 'recipient', '____________________')
    product = field(fields, 'product', '____________________')
    price = field(fields, 'price', '____________________')
    benefits = field(fields, 'benefits', 'высокое качество, гарантия, быстрая доставка')
    contact = field(fields, 'contact', '____________________')
    valid_days = field(fields, 'valid_days', '30')

    heading(doc, 'КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ')
    p = doc.add_paragraph(); p.add_run(f'Дата: {d}').italic = True
    doc.add_paragraph()
    para(doc, f'Уважаемый {recipient}!')
    doc.add_paragraph()
    para(doc, f'Компания {sender} рада предложить Вам сотрудничество и представить наш продукт/услугу:')
    doc.add_paragraph()

    p = doc.add_paragraph(); run = p.add_run(f'{product}'); run.bold = True; run.font.size = Pt(13)
    doc.add_paragraph()

    heading(doc, 'Наши преимущества:', 2)
    for b in [b.strip() for b in benefits.split(',')]:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(b.capitalize())

    doc.add_paragraph()
    heading(doc, 'Стоимость:', 2)
    para(doc, f'Цена: от {price} рублей.')
    para(doc, 'Возможна индивидуальная скидка при объёме заказа.')
    doc.add_paragraph()

    para(doc, f'Предложение действительно в течение {valid_days} дней.')
    doc.add_paragraph()
    para(doc, f'Для связи и уточнения деталей: {contact}')
    doc.add_paragraph()
    para(doc, f'С уважением,\n{sender}')

    return doc, 'Коммерческое_предложение.docx'


# ─── 7. Доверенность ─────────────────────────────────────────────
def gen_power_of_attorney(fields, today):
    doc = base_doc()
    city = field(fields, 'city', 'г. Москва')
    d = field(fields, 'poa_date', today)
    principal = field(fields, 'principal', '____________________')
    principal_passport = field(fields, 'principal_passport', '____________________')
    attorney = field(fields, 'attorney', '____________________')
    attorney_passport = field(fields, 'attorney_passport', '____________________')
    powers = field(fields, 'powers', '____________________')
    valid_until = field(fields, 'valid_until', '____________________')

    heading(doc, 'ДОВЕРЕННОСТЬ')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{city}, {d}')
    doc.add_paragraph()

    para(doc, f'Я, {principal}, паспорт: {principal_passport}, настоящей доверенностью уполномочиваю:')
    doc.add_paragraph()
    para(doc, f'{attorney}, паспорт: {attorney_passport},')
    doc.add_paragraph()
    para(doc, f'совершать от моего имени следующие действия:')
    doc.add_paragraph()
    para(doc, powers)
    doc.add_paragraph()
    para(doc, f'Доверенность выдана сроком до {valid_until}.')
    para(doc, 'Полномочия по настоящей доверенности не могут быть переданы третьим лицам.')
    doc.add_paragraph()
    para(doc, f'Подпись доверителя: _______________  /{principal}/')

    return doc, 'Доверенность.docx'


# ─── 8. NDA (Соглашение о конфиденциальности) ───────────────────
def gen_nda(fields, today):
    doc = base_doc()
    city = field(fields, 'city', 'г. Москва')
    num = field(fields, 'contract_num', '___')
    d = field(fields, 'contract_date', today)
    party1 = field(fields, 'party1', '____________________')
    party2 = field(fields, 'party2', '____________________')
    subject = field(fields, 'subject', 'коммерческая информация, техническая документация, ноу-хау')
    period_years = field(fields, 'period_years', '3')

    heading(doc, 'СОГЛАШЕНИЕ О КОНФИДЕНЦИАЛЬНОСТИ (NDA)')
    heading(doc, f'№ {num}', 2)
    p = doc.add_paragraph(); p.add_run(f'{city}\t\t\t\t\t\t\t\t{d}')
    doc.add_paragraph()
    para(doc, f'{party1} и {party2}, именуемые совместно «Стороны», заключили настоящее соглашение о нижеследующем:')

    heading(doc, '1. ПРЕДМЕТ СОГЛАШЕНИЯ', 2)
    para(doc, f'1.1. Стороны договорились о соблюдении конфиденциальности в отношении следующей информации: {subject}.')
    para(doc, '1.2. Конфиденциальной признаётся любая информация, переданная одной Стороной другой в рамках сотрудничества.')

    heading(doc, '2. ОБЯЗАТЕЛЬСТВА СТОРОН', 2)
    para(doc, '2.1. Каждая из Сторон обязуется не раскрывать конфиденциальную информацию третьим лицам без письменного согласия другой Стороны.')
    para(doc, '2.2. Стороны обязуются использовать конфиденциальную информацию исключительно в целях сотрудничества.')
    para(doc, '2.3. Принимающая Сторона обязана ограничить круг лиц, имеющих доступ к конфиденциальной информации.')

    heading(doc, '3. СРОК ДЕЙСТВИЯ', 2)
    para(doc, f'3.1. Настоящее соглашение вступает в силу с момента подписания и действует в течение {period_years} лет.')
    para(doc, '3.2. Обязательства по конфиденциальности сохраняются в течение 2 лет после истечения срока действия соглашения.')

    heading(doc, '4. ОТВЕТСТВЕННОСТЬ', 2)
    para(doc, '4.1. В случае нарушения настоящего соглашения виновная Сторона обязана возместить причинённые убытки в полном объёме.')

    heading(doc, '5. РЕКВИЗИТЫ И ПОДПИСИ', 2)
    p = doc.add_paragraph(); p.add_run(f'{party1}:\n').bold = True; p.add_run('Подпись: _______________  /_______________/')
    p = doc.add_paragraph(); p.add_run(f'{party2}:\n').bold = True; p.add_run('Подпись: _______________  /_______________/')

    return doc, f'NDA_{num}.docx'


# ─── 9. Приказ по организации ────────────────────────────────────
def gen_order(fields, today):
    doc = base_doc()
    company = field(fields, 'company', '____________________')
    num = field(fields, 'order_num', '___')
    d = field(fields, 'order_date', today)
    city = field(fields, 'city', 'г. Москва')
    subject = field(fields, 'subject', '____________________')
    order_text = field(fields, 'order_text', '____________________')
    responsible = field(fields, 'responsible', '____________________')
    director = field(fields, 'director', '____________________')
    director_position = field(fields, 'director_position', 'Генеральный директор')

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(company.upper())
    run.bold = True

    heading(doc, f'ПРИКАЗ № {num}')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f'{city}\t\t\t\t\t\t\t\t{d}')
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run('О ').bold = False
    run = p.add_run(subject)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    para(doc, 'В целях обеспечения эффективной деятельности организации,')
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('ПРИКАЗЫВАЮ:').bold = True
    doc.add_paragraph()
    para(doc, order_text)
    doc.add_paragraph()
    para(doc, f'Ответственный за исполнение: {responsible}.')
    para(doc, 'Контроль за исполнением настоящего приказа оставляю за собой.')
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.add_run(f'{director_position}      _______________      {director}')

    return doc, f'Приказ_{num}.docx'
